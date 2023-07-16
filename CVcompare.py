import cv2
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
import numpy as np
import os

def find_image_mismatches(reference_image, test_image, output_doc):
    # Load the images
    ref_img = cv2.imread(reference_image)
    test_img = cv2.imread(test_image)

    # Convert images to grayscale
    ref_gray = cv2.cvtColor(ref_img, cv2.COLOR_BGR2GRAY)
    test_gray = cv2.cvtColor(test_img, cv2.COLOR_BGR2GRAY)

    # Compute the absolute difference between the grayscale images
    diff = cv2.absdiff(ref_gray, test_gray)

    # Apply a threshold to the difference image
    _, threshold = cv2.threshold(diff, 30, 255, cv2.THRESH_BINARY)

    # Find contours of the thresholded difference image
    contours, _ = cv2.findContours(threshold, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    # Create a Word document
    doc = Document()
    doc.add_heading('Image Mismatches', level=1)

    # Create a table for displaying the mismatches
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # Set the column widths
    col_widths = [Inches(0.5), Inches(4), Inches(2)]
    for i in range(3):
        table.cell(0, i).width = col_widths[i]

    # Add the header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'No.'
    header_cells[1].text = 'Mismatched Image'
    header_cells[2].text = 'Reference Image'

    # Group neighboring contours within a distance threshold
    grouped_contours = []
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        cx, cy = x + w // 2, y + h // 2

        grouped = False
        for group in grouped_contours:
            if len(group) > 0:
                rotated_rect = cv2.minAreaRect(np.concatenate(group))
                (group_x, group_y), (group_w, group_h), _ = rotated_rect
                group_cx, group_cy = group_x + group_w // 2, group_y + group_h // 2
                distance = np.sqrt((cx - group_cx) ** 2 + (cy - group_cy) ** 2)
                if distance < 50:  # Distance threshold to group contours
                    group.append(contour)
                    x = min(x, group_x)  # Update the leftmost x-coordinate
                    y = min(y, group_y)  # Update the topmost y-coordinate
                    w = max(x + w, group_x + group_w) - x  # Update the total width
                    h = max(y + h, group_y + group_h) - y  # Update the total height
                    grouped = True
                    break

        if not grouped:
            grouped_contours.append([contour])

    # Capture the mismatches and add them to the table
    mismatch_count = 0
    for group in grouped_contours:
        if len(group) > 0:
            # Combine contours in the group into a single contour
            combined_contour = np.concatenate(group)

            # Find the bounding box for the combined contour
            x, y, w, h = cv2.boundingRect(combined_contour)
            margin = 50  # Margin around the error region
            x -= margin
            y -= margin
            w += 2 * margin
            h += 2 * margin

            # Capture the screenshot of the mismatch with margin
            mismatch_img = test_img[max(0, y):min(y + h, test_img.shape[0]), max(0, x):min(x + w, test_img.shape[1])]

            # Highlight the error region in red
            cv2.rectangle(mismatch_img, (max(0, margin), max(0, margin)), (min(w, w - margin), min(h, h - margin)), (0, 0, 255), 2)

            # Capture the corresponding part of the original reference image
            ref_mismatch_img = ref_img[max(0, y):min(y + h, ref_img.shape[0]), max(0, x):min(x + w, ref_img.shape[1])]

            # Save the screenshots
            mismatch_filename = f"mismatch_{mismatch_count + 1}.png"
            ref_mismatch_filename = f"ref_mismatch_{mismatch_count + 1}.png"
            cv2.imwrite(mismatch_filename, mismatch_img)
            cv2.imwrite(ref_mismatch_filename, ref_mismatch_img)

            # Add the mismatch details and images to the table
            row_cells = table.add_row().cells
            row_cells[0].text = str(mismatch_count + 1)
            row_cells[1].add_paragraph().add_run().add_picture(mismatch_filename, width=Inches(3), height=Inches(2))
            row_cells[2].add_paragraph().add_run().add_picture(ref_mismatch_filename, width=Inches(1.5), height=Inches(1.5))

            # Set the vertical alignment of cells
            for cell in row_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            mismatch_count += 1

    # Save and close the Word document
    doc.save(output_doc)

    # Delete the mismatched image files after they have been used
    for i in range(mismatch_count):
        mismatch_filename = f"mismatch_{i + 1}.png"
        ref_mismatch_filename = f"ref_mismatch_{i + 1}.png"
        os.remove(mismatch_filename)
        os.remove(ref_mismatch_filename)

# Test the function
reference_image_path = 'PSDMobile.png'
test_image_path = 'croppedMobile.png'
output_document = 'mismatches.docx'
find_image_mismatches(reference_image_path, test_image_path, output_document)

